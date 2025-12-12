import io
import json
import requests
import ast
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

def get_doubao_response(user_info, form_context):
    # 确保输入是 UTF-8 编码的字符串
    if isinstance(user_info, bytes):
        user_info = user_info.decode('utf-8')
    if isinstance(form_context, bytes):
        form_context = form_context.decode('utf-8')

    # 设置API密钥和模型端点 - 可以从环境变量或硬编码获取
    url = "https://ark.cn-beijing.volces.com/api/v3/chat/completions"
    api_key = os.environ.get("ARK_API_KEY") or "5410d463-1115-4320-9279-a5441ce30694"
    model_endpoint = os.environ.get("MODEL_ENDPOINT") or "doubao-seed-1-6-251015"

    # 打印调试信息
    print(f"🔑 API Key: {api_key[:10]}...")
    print(f"🤖 Model: {model_endpoint}")
    print(f"📝 User Info Length: {len(user_info)}")
    print(f"📋 Form Context Length: {len(form_context)}")

    prompt = f"""
    你是一个智能填表助手。
    【任务】
    表格中的空缺项已标记为 {{1}}, {{2}}...
    请根据【个人资料】推断内容。
    【个人资料】
    {user_info}
    【表格上下文】
    {form_context}
    【要求】
    1. 返回纯 JSON，格式 {{"{{1}}": "内容"}}。
    2. 找不到信息填 "无"。
    """
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {"model": model_endpoint, "messages": [{"role": "user", "content": prompt}], "temperature": 0.1}
    try:
        response = requests.post(url, headers=headers, json=data)
        print(f"📡 Response Status: {response.status_code}")
        if response.status_code != 200:
            print(f"❌ Error Response: {response.text}")
            return {}

        res_json = response.json()
        print(f"🔍 Response Keys: {res_json.keys()}")

        if 'choices' not in res_json or not res_json['choices']:
            print("❌ No choices in response")
            return {}

        content = res_json['choices'][0]['message']['content'].replace("```json", "").replace("```", "").strip()
        print(f"📄 Raw Content: {content[:200]}...")

        try:
            result = json.loads(content)
            print(f"✅ Parsed JSON: {result}")
            return result
        except json.JSONDecodeError as e:
            print(f"⚠️ JSON decode failed: {e}")
            try:
                result = ast.literal_eval(content)
                print(f"✅ Parsed with ast: {result}")
                return result
            except Exception as e2:
                print(f"❌ AST parse failed: {e2}")
                return {}
    except Exception as e:
        print(f"❌ Exception: {e}")
        import traceback
        traceback.print_exc()
        return {}

def fill_form(docx_bytes, user_info_text, photo_bytes):
    doc = Document(io.BytesIO(docx_bytes))
    photo_coords = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text_lower = cell.text.lower()
                if ("照片" in text_lower) or ("相片" in text_lower) or ("证件照" in text_lower):
                    photo_coords.append((t_idx, r_idx, c_idx))
    if photo_coords and photo_bytes:
        for (t_idx, r_idx, c_idx) in photo_coords:
            cell = doc.tables[t_idx].rows[r_idx].cells[c_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(io.BytesIO(photo_bytes), width=Cm(3.5))
    counter = 1
    placeholder_map = {}
    form_context_lines = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            row_text_parts = []
            has_placeholder = False
            for c_idx, cell in enumerate(row.cells):
                if (t_idx, r_idx, c_idx) in photo_coords:
                    row_text_parts.append("[照片]")
                    continue
                text = cell.text.strip()
                if not text:
                    placeholder = f"{{{counter}}}"
                    cell.text = placeholder
                    placeholder_map[placeholder] = cell
                    row_text_parts.append(placeholder)
                    has_placeholder = True
                    counter += 1
                else:
                    row_text_parts.append(text)
            if has_placeholder:
                form_context_lines.append(" | ".join(row_text_parts))
    if not placeholder_map:
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    fill_data = get_doubao_response(user_info_text, "\n".join(form_context_lines))
    if fill_data:
        for key, value in fill_data.items():
            target_key = key if key.startswith("{") else f"{{{key}}}"
            if target_key in placeholder_map:
                cell = placeholder_map[target_key]
                cell.text = str(value)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

